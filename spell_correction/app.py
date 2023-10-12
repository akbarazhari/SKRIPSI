from flask import Flask, render_template, request, send_file
import re
import docx
import json
import os
from collections import Counter
from werkzeug.utils import secure_filename

app = Flask(__name__)

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['HASIL_FOLDER'] = 'hasil'

def words(text): return re.findall(r'\w+', text.lower())

WORDS = Counter(words(open('katadasar.txt').read()))

def P(word, N=sum(WORDS.values())):
    # "Probability of `word`."
    return WORDS[word] / N

def correction(word):
    # "Most probable spelling correction for word."
    return max(candidates(word), key=P)

def candidates(word):
    # "Generate possible spelling corrections for word."
    return (known([word]) or known(edits1(word)) or known(edits2(word)) or [word])

def known(words):
    # "The subset of `words` that appear in the dictionary of WORDS."
    return set(w for w in words if w in WORDS)

def edits1(word): 
    # "All edits that are one edit away from `word`."
    letters    = 'abcdefghijklmnopqrstuvwxyz'
    splits     = [(word[:i], word[i:])    for i in range(len(word) + 1)] # [('', 'kemarin'), ('k', 'emarin'), ('ke', 'marin'), dst]
    deletes    = [L + R[1:]               for L, R in splits if R] # ['emarin', 'kmarin', 'kearin', dst]
    transposes = [L + R[1] + R[0] + R[2:] for L, R in splits if len(R)>1] # ['ekmarin', 'kmearin', 'keamrin', dst]
    replaces   = [L + c + R[1:]           for L, R in splits if R for c in letters] # ['aemarin', 'bemarin', 'cemarin', dst]
    inserts    = [L + c + R               for L, R in splits for c in letters] # ['akemarin', 'bkemarin', 'ckemarin', dst]
    return set(deletes + transposes + replaces + inserts)

def edits2(word):
    # "All edits that are two edits away from `word`."
    return (e2 for e1 in edits1(word) for e2 in edits1(e1))

# LAVENSTHEIN DISTANCE

def edit_distance(str1, str2):
    m = len(str1) + 1
    n = len(str2) + 1
    dp = [[0] * n for _ in range(m)]
    for i in range(m):
        dp[i][0] = i
    for j in range(n):
        dp[0][j] = j
    for i in range(1, m):
        for j in range(1, n):
            if str1[i-1] == str2[j-1]:
                dp[i][j] = dp[i-1][j-1]
            else:
                dp[i][j] = min(dp[i][j-1], dp[i-1][j], dp[i-1][j-1]) + 1
    return dp[m-1][n-1]

def correct_spelling(hasil, WORDS):
    corrected_text = ""
    words = hasil.split()
    for word in words:
        if word in WORDS:
            corrected_text += word + " "
            continue
        min_dist = float("inf")
        best_word = word
        for w in WORDS:
            dist = edit_distance(word, w)
            if dist < min_dist:
                min_dist = dist
                best_word = w
        corrected_text += best_word + " "
    return corrected_text

@app.route('/', methods=['GET'])
def index():
    # Main page
    return render_template('index.html')

@app.route('/predict', methods=['GET', 'POST'])
def predict():
    if request.method == 'POST':
        f = request.files['file']
        basepath = os.path.dirname(__file__)
        file_path = os.path.join(
            basepath, 'uploads', secure_filename(f.filename))
    
        f.save(file_path)
        docPaths = ''
        fixDocPath = docPaths + file_path
        doc = docx.Document(fixDocPath)

        for para in doc.paragraphs:
            for run in para.runs:
                corrected_para = []
                if run.text:
                    for word in para.text.split():
                        corrected_word = correction(word)
                        corrected_para.append(corrected_word)
                    para.text = " ".join(corrected_para)
        hasil = json.dumps(corrected_para)
        corrected_text = correct_spelling(hasil, WORDS)

        doc.save('corrected_document.docx')

        doc2 = docx.Document("new.docx")
        list_kata_salah = []
        list_kata_benar = []
        for para in doc2.paragraphs:
            words = para.text.split()
            for i in words:
                if i not in WORDS:
                    list_kata_salah.append(i)
                    continue
                else:
                    list_kata_benar.append(i)
                    continue

        rsl = []
        for c in list_kata_salah:
            rsl.append(f"'{c}' -> '{correct_spelling(c, WORDS)}'")

        return render_template('index.html', teks=corrected_text, benar=rsl, salah=list_kata_salah)

@app.route('/download', methods=['GET', 'POST'])
def download():

    # mengambil path file yang telah di proses
    output_file = os.path.abspath('corrected_document.docx')

    # mengirimkan file hasil proses ke user
    return send_file(output_file, as_attachment=True)
    
if __name__ == '__main__':
        app.run(debug=True)